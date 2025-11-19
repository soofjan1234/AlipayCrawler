"""
数据导出模块
用于将提取的B站数据导出为Excel和Word格式
"""

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import os
from typing import List, Dict, Any
from loguru import logger


class DataExporter:
    """数据导出器"""
    
    def __init__(self, output_dir: str = "/Users/Zhuanz/projects/PythonWS/Alipay/data"):
        """
        初始化数据导出器
        
        Args:
            output_dir: 输出目录路径
        """
        self.output_dir = output_dir
        self.bilibili_data_dir = os.path.join(output_dir, "bilibili_data")
        
        # 确保输出目录存在
        os.makedirs(self.bilibili_data_dir, exist_ok=True)
        
    def export_to_excel(self, contents_data: List[Dict[str, Any]], filename: str = "bilibili_data.xlsx") -> str:
        """
        将数据导出为Excel格式
        
        Args:
            contents_data: 提取的内容数据列表
            filename: 输出文件名
            
        Returns:
            str: 输出文件的完整路径
        """
        try:
            # 准备Excel数据
            excel_data = []
            
            for content in contents_data:
                # 提取图片链接（如果有的话）
                image_urls = content.get("图片链接", "")
                
                # 确定内容类型
                content_type = content.get("内容类型", "动态")
                if content_type == "视频":
                    content_type = "视频"
                else:
                    content_type = "图文"
                
                # 解析发布时间
                publish_time = content.get("发布时间_解析", "")
                if publish_time:
                    try:
                        # 尝试转换为datetime格式
                        publish_dt = datetime.strptime(publish_time, "%Y-%m-%d")
                        publish_time = publish_dt
                    except ValueError:
                        publish_time = ""
                
                # 构建Excel行数据
                excel_row = {
                    "content_id": content.get("内容ID", ""),
                    "content_type": content_type,
                    "text_content": content.get("文案内容", ""),
                    "publish_time": publish_time,
                    "like_count": self._extract_number(content.get("点赞数", "0")),
                    "comment_count": self._extract_number(content.get("评论数", "0")),
                    "repost_count": self._extract_number(content.get("转发数", "0")),
                    "image_urls": image_urls,
                    "platform": "bilibili"
                }
                
                excel_data.append(excel_row)
            
            # 创建DataFrame
            df = pd.DataFrame(excel_data)
            
            # 按发布时间倒序排列
            if not df.empty and 'publish_time' in df.columns:
                df = df.sort_values('publish_time', ascending=False)
            
            # 保存到Excel文件
            output_path = os.path.join(self.bilibili_data_dir, filename)
            df.to_excel(output_path, index=False, engine='openpyxl')
            
            logger.info(f"✅ Excel文件已导出: {output_path}")
            logger.info(f"📊 共导出 {len(excel_data)} 条数据")
            
            return output_path
            
        except Exception as e:
            logger.error(f"导出Excel文件时发生错误: {str(e)}")
            raise
    
    def export_to_word(self, contents_data: List[Dict[str, Any]], filename: str = "bilibili_content.docx") -> str:
        """
        将数据导出为Word格式，按月份分组
        
        Args:
            contents_data: 提取的内容数据列表
            filename: 输出文件名
            
        Returns:
            str: 输出文件的完整路径
        """
        try:
            # 创建Word文档
            doc = Document()
            
            # 设置标题
            title = doc.add_heading('支付宝B站动态内容汇总（2024年5月1日 - 2024年11月1日）', 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 按月份分组数据
            monthly_data = self._group_by_month(contents_data)
            
            # 按月份倒序排列（11月到5月）
            months = ["2024年11月", "2024年10月", "2024年9月", "2024年8月", "2024年7月", "2024年6月", "2024年5月"]
            
            for month in months:
                if month in monthly_data:
                    # 添加月份标题
                    month_heading = doc.add_heading(f'=== {month} ===', level=1)
                    month_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    
                    # 按时间倒序排列该月的内容
                    month_contents = sorted(monthly_data[month], 
                                          key=lambda x: x.get("发布时间_解析", ""), 
                                          reverse=True)
                    
                    # 添加该月的每条内容
                    for content in month_contents:
                        # 发布时间
                        publish_time = content.get("发布时间_原始", "")
                        if publish_time:
                            time_para = doc.add_paragraph()
                            time_run = time_para.add_run(f"【发布时间】{publish_time}")
                            time_run.font.size = Pt(11)
                            time_run.bold = True
                        
                        # 文案内容
                        text_content = content.get("文案内容", "")
                        if text_content:
                            text_para = doc.add_paragraph()
                            text_run = text_para.add_run(f"【文案内容】{text_content}")
                            text_run.font.size = Pt(11)
                        
                        # 添加统计数据
                        like_count = self._extract_number(content.get("点赞数", "0"))
                        comment_count = self._extract_number(content.get("评论数", "0"))
                        repost_count = self._extract_number(content.get("转发数", "0"))
                        
                        # 统计信息
                        stats_para = doc.add_paragraph()
                        stats_run = stats_para.add_run(f"【统计数据】点赞：{like_count} | 评论：{comment_count} | 转发：{repost_count}")
                        stats_run.font.size = Pt(10)
                        stats_run.font.color.rgb = RGBColor(128, 128, 128)  # 灰色字体
                        
                        # 如果是视频，添加视频描述
                        if content.get("内容类型") == "视频":
                            video_desc = content.get("视频描述", "")
                            if video_desc:
                                video_para = doc.add_paragraph()
                                video_run = video_para.add_run(f"【视频描述】{video_desc}")
                                video_run.font.size = Pt(11)
                                video_run.italic = True
                        
                        # 添加空行分隔
                        doc.add_paragraph()
            
            # 保存Word文档
            output_path = os.path.join(self.bilibili_data_dir, filename)
            doc.save(output_path)
            
            logger.info(f"✅ Word文档已导出: {output_path}")
            logger.info(f"📄 共导出 {len(contents_data)} 条内容")
            
            return output_path
            
        except Exception as e:
            logger.error(f"导出Word文档时发生错误: {str(e)}")
            raise
    
    def _group_by_month(self, contents_data: List[Dict[str, Any]]) -> Dict[str, List[Dict[str, Any]]]:
        """
        按月份分组数据
        
        Args:
            contents_data: 内容数据列表
            
        Returns:
            Dict: 按月份分组的数据字典
        """
        monthly_data = {}
        
        for content in contents_data:
            publish_date = content.get("发布时间_解析", "")
            if publish_date:
                try:
                    # 解析日期
                    date_obj = datetime.strptime(publish_date, "%Y-%m-%d")
                    month_key = f"{date_obj.year}年{date_obj.month}月"
                    
                    if month_key not in monthly_data:
                        monthly_data[month_key] = []
                    
                    monthly_data[month_key].append(content)
                    
                except ValueError:
                    logger.warning(f"无法解析日期: {publish_date}")
                    continue
        
        return monthly_data
    
    def _extract_number(self, text: str) -> int:
        """
        从文本中提取数字
        
        Args:
            text: 包含数字的文本
            
        Returns:
            int: 提取的数字
        """
        try:
            # 移除非数字字符，保留数字
            import re
            numbers = re.findall(r'\d+', str(text))
            return int(numbers[0]) if numbers else 0
        except Exception:
            return 0
    
    def export_all_formats(self, contents_data: List[Dict[str, Any]]) -> Dict[str, str]:
        """
        导出所有格式的文件
        
        Args:
            contents_data: 内容数据列表
            
        Returns:
            Dict: 包含各文件路径的字典
        """
        results = {}
        
        try:
            # 导出Excel
            excel_path = self.export_to_excel(contents_data)
            results["excel"] = excel_path
            
            # 导出Word
            word_path = self.export_to_word(contents_data)
            results["word"] = word_path
            
            logger.info("✅ 所有格式文件导出完成")
            
        except Exception as e:
            logger.error(f"导出文件时发生错误: {str(e)}")
            raise
        
        return results