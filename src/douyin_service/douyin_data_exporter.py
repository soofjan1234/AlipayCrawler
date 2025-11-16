"""
抖音数据导出模块
用于将提取的抖音数据导出为Excel和Word格式
"""

import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import os
import re
from typing import List, Dict, Any
from loguru import logger


class DouyinDataExporter:
    """抖音数据导出器"""
    
    def __init__(self, output_dir: str = "/Users/Zhuanz/projects/PythonWS/Alipay/data"):
        """
        初始化数据导出器
        
        Args:
            output_dir: 输出目录路径
        """
        self.output_dir = output_dir
        self.douyin_data_dir = os.path.join(output_dir, "douyin_data")
        
        # 确保输出目录存在
        os.makedirs(self.douyin_data_dir, exist_ok=True)
        
    def parse_douyin_data(self, stats_file: str = "/Users/Zhuanz/projects/PythonWS/Alipay/3.txt", 
                         content_file: str = "/Users/Zhuanz/projects/PythonWS/Alipay/2.txt") -> List[Dict[str, Any]]:
        """
        解析抖音数据文件，整合统计数据和文案内容
        
        Args:
            stats_file: 统计数据文件路径（3.txt）
            content_file: 文案内容文件路径（2.txt）
            
        Returns:
            List[Dict]: 整合后的数据列表
        """
        try:
            # 读取统计数据文件
            stats_data = self._parse_stats_file(stats_file)
            
            # 读取文案内容文件
            content_data = self._parse_content_file(content_file)
            
            # 整合数据
            merged_data = []
            for stats in stats_data:
                video_url = stats.get("video_url", "")
                # 从文案文件中找到对应的内容
                content = self._find_content_by_url(video_url, content_data)
                
                merged_item = {
                    "video_url": video_url,
                    "content_text": content.get("content_text", ""),
                    "publish_time": stats.get("publish_time", ""),
                    "like_count": stats.get("like_count", 0),
                    "comment_count": stats.get("comment_count", 0),
                    "collect_count": stats.get("collect_count", 0),
                    "share_count": stats.get("share_count", 0),
                    "publish_time_parsed": self._parse_publish_time(stats.get("publish_time", ""))
                }
                
                merged_data.append(merged_item)
            
            logger.info(f"✅ 成功解析抖音数据，共 {len(merged_data)} 条记录")
            return merged_data
            
        except Exception as e:
            logger.error(f"解析抖音数据时发生错误: {str(e)}")
            raise
    
    def _parse_stats_file(self, file_path: str) -> List[Dict[str, Any]]:
        """解析统计数据文件（3.txt）"""
        stats_data = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # 按视频分割内容
            video_sections = re.split(r'=== 视频URL: (.*?) ===', content)
            
            for i in range(1, len(video_sections), 2):
                if i + 1 < len(video_sections):
                    url = video_sections[i].strip()
                    section_content = video_sections[i + 1]
                    
                    # 提取各项数据
                    stats = {
                        "video_url": url,
                        "like_count": self._extract_number_from_text(section_content, "点赞数"),
                        "comment_count": self._extract_number_from_text(section_content, "评论数"),
                        "collect_count": self._extract_number_from_text(section_content, "收藏数"),
                        "share_count": self._extract_number_from_text(section_content, "转发数"),
                        "publish_time": self._extract_publish_time(section_content)
                    }
                    
                    stats_data.append(stats)
            
            logger.info(f"从统计数据文件中解析出 {len(stats_data)} 条记录")
            
        except Exception as e:
            logger.error(f"解析统计数据文件时发生错误: {str(e)}")
            
        return stats_data
    
    def _parse_content_file(self, file_path: str) -> List[Dict[str, Any]]:
        """解析文案内容文件（2.txt）"""
        content_data = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                lines = f.readlines()
            
            current_url = ""
            current_content = []
            
            for line in lines:
                line = line.strip()
                if line.startswith('https://www.douyin.com/video/'):
                    # 保存前一个视频的内容
                    if current_url and current_content:
                        content_data.append({
                            "video_url": current_url,
                            "content_text": "\n".join(current_content).strip()
                        })
                    
                    # 开始新的视频
                    current_url = line
                    current_content = []
                elif line and current_url:
                    current_content.append(line)
            
            # 保存最后一个视频的内容
            if current_url and current_content:
                content_data.append({
                    "video_url": current_url,
                    "content_text": "\n".join(current_content).strip()
                })
            
            logger.info(f"从文案内容文件中解析出 {len(content_data)} 条记录")
            
        except Exception as e:
            logger.error(f"解析文案内容文件时发生错误: {str(e)}")
            
        return content_data
    
    def _find_content_by_url(self, url: str, content_data: List[Dict[str, Any]]) -> Dict[str, Any]:
        """根据URL查找对应的文案内容"""
        for content in content_data:
            if content.get("video_url", "") == url:
                return content
        return {"content_text": ""}
    
    def _extract_number_from_text(self, text: str, field_name: str) -> int:
        """从文本中提取指定字段的数值"""
        try:
            pattern = f"{field_name}: (\\d+)"
            match = re.search(pattern, text)
            if match:
                return int(match.group(1))
        except Exception:
            pass
        return 0
    
    def _extract_publish_time(self, text: str) -> str:
        """从文本中提取发布时间"""
        try:
            pattern = r"发布时间: (.*?)\n"
            match = re.search(pattern, text)
            if match:
                return match.group(1).strip()
        except Exception:
            pass
        return ""
    
    def _parse_publish_time(self, publish_time: str) -> str:
        """解析发布时间为标准格式"""
        try:
            # 处理格式如 "发布时间：2025-10-18 09:01"
            if "发布时间：" in publish_time:
                time_str = publish_time.replace("发布时间：", "").strip()
                # 尝试解析为datetime对象
                dt = datetime.strptime(time_str, "%Y-%m-%d %H:%M")
                return dt.strftime("%Y-%m-%d")
        except Exception:
            pass
        return ""
    
    def export_to_excel(self, data: List[Dict[str, Any]], filename: str = "douyin_data.xlsx") -> str:
        """
        将数据导出为Excel格式
        
        Args:
            data: 要导出的数据列表
            filename: 输出文件名
            
        Returns:
            str: 输出文件的完整路径
        """
        try:
            # 准备Excel数据
            excel_data = []
            
            for item in data:
                excel_row = {
                    "视频URL": item.get("video_url", ""),
                    "文案内容": item.get("content_text", ""),
                    "发布时间": item.get("publish_time_parsed", ""),
                    "点赞数": item.get("like_count", 0),
                    "评论数": item.get("comment_count", 0),
                    "收藏数": item.get("collect_count", 0),
                    "转发数": item.get("share_count", 0),
                    "平台": "抖音"
                }
                
                excel_data.append(excel_row)
            
            # 创建DataFrame
            df = pd.DataFrame(excel_data)
            
            # 按发布时间倒序排列
            if not df.empty and '发布时间' in df.columns:
                df = df.sort_values('发布时间', ascending=False)
            
            # 保存到Excel文件
            output_path = os.path.join(self.douyin_data_dir, filename)
            df.to_excel(output_path, index=False, engine='openpyxl')
            
            logger.info(f"✅ Excel文件已导出: {output_path}")
            logger.info(f"📊 共导出 {len(excel_data)} 条数据")
            
            return output_path
            
        except Exception as e:
            logger.error(f"导出Excel文件时发生错误: {str(e)}")
            raise
    
    def export_to_word(self, data: List[Dict[str, Any]], filename: str = "douyin_content.docx") -> str:
        """
        将数据导出为Word格式，按顺序摆放，只保留时间和内容
        
        Args:
            data: 要导出的数据列表
            filename: 输出文件名
            
        Returns:
            str: 输出文件的完整路径
        """
        try:
            # 创建Word文档
            doc = Document()
            
            # 设置标题
            title = doc.add_heading('支付宝抖音视频内容汇总', 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 按发布时间倒序排列数据
            sorted_data = sorted(data, 
                              key=lambda x: x.get("publish_time_parsed", ""), 
                              reverse=True)
            
            # 添加每条内容（只保留时间和内容）
            for item in sorted_data:
                # 发布时间
                publish_time = item.get("publish_time", "")
                if publish_time:
                    time_para = doc.add_paragraph()
                    time_run = time_para.add_run(f"【发布时间】{publish_time}")
                    time_run.font.size = Pt(11)
                    time_run.bold = True
                
                # 文案内容
                content_text = item.get("content_text", "")
                if content_text:
                    text_para = doc.add_paragraph()
                    text_run = text_para.add_run(f"【文案内容】{content_text}")
                    text_run.font.size = Pt(11)
                
                # 添加空行分隔
                doc.add_paragraph()
            
            # 保存Word文档
            output_path = os.path.join(self.douyin_data_dir, filename)
            doc.save(output_path)
            
            logger.info(f"✅ Word文档已导出: {output_path}")
            logger.info(f"📄 共导出 {len(data)} 条内容")
            
            return output_path
            
        except Exception as e:
            logger.error(f"导出Word文档时发生错误: {str(e)}")
            raise
    
    def _group_by_month(self, data: List[Dict[str, Any]]) -> Dict[str, List[Dict[str, Any]]]:
        """
        按月份分组数据
        
        Args:
            data: 数据列表
            
        Returns:
            Dict: 按月份分组的数据字典
        """
        monthly_data = {}
        
        for item in data:
            publish_date = item.get("publish_time_parsed", "")
            if publish_date:
                try:
                    # 解析日期
                    date_obj = datetime.strptime(publish_date, "%Y-%m-%d")
                    month_key = f"{date_obj.year}年{date_obj.month}月"
                    
                    if month_key not in monthly_data:
                        monthly_data[month_key] = []
                    
                    monthly_data[month_key].append(item)
                    
                except ValueError:
                    logger.warning(f"无法解析日期: {publish_date}")
                    continue
            else:
                # 没有日期的数据放入"未知时间"分组
                if "未知时间" not in monthly_data:
                    monthly_data["未知时间"] = []
                monthly_data["未知时间"].append(item)
        
        return monthly_data
    
    def export_all_formats(self, data: List[Dict[str, Any]]) -> Dict[str, str]:
        """
        导出所有格式的文件
        
        Args:
            data: 要导出的数据列表
            
        Returns:
            Dict: 包含各文件路径的字典
        """
        results = {}
        
        try:
            # 导出Excel
            excel_path = self.export_to_excel(data)
            results["excel"] = excel_path
            
            # 导出Word
            word_path = self.export_to_word(data)
            results["word"] = word_path
            
            logger.info("✅ 所有格式文件导出完成")
            
        except Exception as e:
            logger.error(f"导出文件时发生错误: {str(e)}")
            raise
        
        return results


# 需要导入RGBColor用于设置字体颜色
from docx.shared import RGBColor