#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import re
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import os

class BilibiliDataExporter:
    """B站数据导出器"""
    
    def __init__(self, txt_file_path: str):
        self.txt_file_path = txt_file_path
        self.data = []
        
    def parse_txt_data(self):
        """解析1.txt文件中的数据"""
        with open(self.txt_file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 使用正则表达式解析每个内容块
        content_blocks = re.split(r'内容 \d+:', content)
        
        for block in content_blocks:
            if not block.strip():
                continue
                
            # 解析各个字段
            data_item = {}
            
            # 内容ID
            content_id_match = re.search(r'内容ID: (\d+)', block)
            if content_id_match:
                data_item['content_id'] = content_id_match.group(1)
            
            # 作者
            author_match = re.search(r'作者: (.+)', block)
            if author_match:
                data_item['author'] = author_match.group(1)
            
            # 发布时间
            publish_time_match = re.search(r'发布时间: (.+)', block)
            if publish_time_match:
                raw_time = publish_time_match.group(1)
                # 转换为标准格式
                data_item['publish_time_raw'] = raw_time
                data_item['publish_time'] = self._parse_publish_time(raw_time)
            
            # 文案内容
            text_content_match = re.search(r'文案内容: (.+?)(?=\s*内容类型:)', block, re.DOTALL)
            if text_content_match:
                data_item['text_content'] = text_content_match.group(1).strip()
            
            # 内容类型
            content_type_match = re.search(r'内容类型: (.+)', block)
            if content_type_match:
                data_item['content_type'] = content_type_match.group(1)
            
            # 视频描述（如果存在）
            video_desc_match = re.search(r'视频描述: (.+?)(?=\s*点赞数:)', block, re.DOTALL)
            if video_desc_match:
                data_item['video_description'] = video_desc_match.group(1).strip()
            
            # 点赞数
            like_match = re.search(r'点赞数: (\d+)', block)
            if like_match:
                data_item['like_count'] = int(like_match.group(1))
            
            # 评论数
            comment_match = re.search(r'评论数: (\d+)', block)
            if comment_match:
                data_item['comment_count'] = int(comment_match.group(1))
            
            # 转发数
            repost_match = re.search(r'转发数: (\d+)', block)
            if repost_match:
                data_item['repost_count'] = int(repost_match.group(1))
            
            # 图片链接
            image_link_match = re.search(r'图片链接: (.+?)(?=\s*视频链接:|\s*卡片高度:|\s*提取轮次:|$)', block)
            if image_link_match:
                data_item['image_link'] = image_link_match.group(1).strip()
            
            # 视频链接
            video_link_match = re.search(r'视频链接: (.+?)(?=\s*卡片高度:|\s*提取轮次:|$)', block)
            if video_link_match:
                data_item['video_link'] = video_link_match.group(1).strip()
            
            # 平台标识
            data_item['platform'] = 'bilibili'
            
            if data_item.get('content_id'):
                self.data.append(data_item)
        
        print(f"成功解析 {len(self.data)} 条数据")
        return self.data
    
    def _parse_publish_time(self, time_str: str) -> str:
        """解析发布时间为标准格式"""
        try:
            # 处理 "MM月DD日" 格式
            if '月' in time_str and '日' in time_str:
                # 假设是2025年的数据
                month_day = time_str.replace(' · 投稿了视频', '').strip()
                month = int(re.search(r'(\d+)月', month_day).group(1))
                day = int(re.search(r'(\d+)日', month_day).group(1))
                return f"2025-{month:02d}-{day:02d}"
        except:
            pass
        
        return time_str
    
    def export_to_excel(self, output_path: str = 'bilibili_data.xlsx'):
        """导出到Excel文件"""
        if not self.data:
            self.parse_txt_data()
        
        # 创建DataFrame
        df_data = []
        for item in self.data:
            row = {
                '内容ID': item.get('content_id', ''),
                '内容类型': item.get('content_type', ''),
                '文案内容': item.get('text_content', ''),
                '发布时间': item.get('publish_time', ''),
                '点赞数': item.get('like_count', 0),
                '评论数': item.get('comment_count', 0),
                '转发数': item.get('repost_count', 0),
                '图片链接': item.get('image_link', ''),
                '平台标识': item.get('platform', 'bilibili')
            }
            df_data.append(row)
        
        df = pd.DataFrame(df_data)
        
        # 导出到Excel
        df.to_excel(output_path, index=False, engine='openpyxl')
        print(f"Excel文件已导出: {output_path}")
        return output_path
    
    def export_to_word(self, output_path: str = 'bilibili_content.docx'):
        """导出到Word文件"""
        if not self.data:
            self.parse_txt_data()
        
        # 创建Word文档
        doc = Document()
        
        # 设置标题
        title = doc.add_heading('支付宝B站动态内容汇总（2025年5月1日 - 2025年11月1日）', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 直接按1.txt的原始顺序添加内容
        for item in self.data:
            # 发布时间
            time_para = doc.add_paragraph()
            time_para.add_run('【发布时间】').bold = True
            time_para.add_run(f' {item.get("publish_time", "")}')
            
            # 文案内容
            content_para = doc.add_paragraph()
            content_para.add_run('【文案内容】').bold = True
            content_para.add_run(f' {item.get("text_content", "")}')
            
            # 如果有视频链接，添加视频链接
            if item.get('video_link'):
                video_para = doc.add_paragraph()
                video_para.add_run('【视频链接】').bold = True
                video_para.add_run(f' {item.get("video_link", "")}')
            
            # 如果是视频类型，添加视频描述
            if item.get('content_type') == '视频' and item.get('video_description'):
                video_desc_para = doc.add_paragraph()
                video_desc_para.add_run('【视频描述】').bold = True
                video_desc_para.add_run(f' {item.get("video_description", "")}')
            
            # 添加空行分隔
            doc.add_paragraph()
        
        # 保存文档
        doc.save(output_path)
        print(f"Word文件已导出: {output_path}")
        return output_path
    
    def _get_month(self, publish_time: str) -> str:
        """从发布时间中提取月份"""
        try:
            if '-' in publish_time:
                parts = publish_time.split('-')
                if len(parts) >= 2:
                    year = parts[0]
                    month = parts[1]
                    return f"{year}年{int(month)}月"
        except:
            pass
        return '未知月份'
    
    def export_all(self):
        """导出所有格式"""
        excel_path = self.export_to_excel()
        word_path = self.export_to_word()
        return excel_path, word_path

if __name__ == "__main__":
    # 使用示例
    exporter = BilibiliDataExporter('/Users/Zhuanz/projects/PythonWS/Alipay/1.txt')
    exporter.export_all()