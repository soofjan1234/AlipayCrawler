#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document

def check_word_content():
    """检查Word文档内容"""
    try:
        doc = Document('bilibili_content.docx')
        print(f"Word文档段落数: {len(doc.paragraphs)}")
        
        # 检查前几个内容项的段落
        image_count = 0
        video_count = 0
        
        for i, para in enumerate(doc.paragraphs[:50]):  # 只检查前50个段落
            text = para.text.strip()
            if '图片链接' in text:
                image_count += 1
                print(f"发现图片链接段落 {i+1}: {text[:100]}...")
            if '视频链接' in text:
                video_count += 1
                print(f"发现视频链接段落 {i+1}: {text[:100]}...")
        
        print(f"\n图片链接段落数量: {image_count}")
        print(f"视频链接段落数量: {video_count}")
        
        return True
    except Exception as e:
        print(f"检查Word文件时出错: {e}")
        return False

if __name__ == "__main__":
    check_word_content()